"""
GhostPC Document Module — Flagship Feature
Excel read/write, PDF generation, AI-powered report creation, DOCX, form filling.
"""

import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Optional, Union

logger = logging.getLogger(__name__)


def _temp_dir() -> Path:
    from config import TEMP_DIR
    return TEMP_DIR


# ─── Excel Operations ────────────────────────────────────────────────────────

def read_excel(path: str) -> dict:
    """
    Read all sheets from an Excel file.
    Returns structured data as dict: { sheet_name: [ {col: val, ...}, ... ] }
    """
    try:
        import openpyxl

        p = Path(path).expanduser()
        if not p.exists():
            return {"success": False, "error": f"File not found: {path}"}

        wb = openpyxl.load_workbook(str(p), data_only=True)
        result = {}

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                result[sheet_name] = []
                continue

            # First row as headers
            headers = [str(h) if h is not None else f"Col{i}" for i, h in enumerate(rows[0])]
            data = []
            for row in rows[1:]:
                if any(cell is not None for cell in row):
                    record = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
                    data.append(record)
            result[sheet_name] = data

        # Build summary text
        total_rows = sum(len(v) for v in result.values())
        sheets_summary = ", ".join(f"{k} ({len(v)} rows)" for k, v in result.items())
        text = f"📊 Excel: {p.name}\nSheets: {sheets_summary}\nTotal rows: {total_rows}"

        return {
            "success": True,
            "data": result,
            "sheets": list(result.keys()),
            "total_rows": total_rows,
            "text": text,
            "file_name": p.name,
        }

    except Exception as e:
        logger.error(f"read_excel error: {e}")
        return {"success": False, "error": str(e)}


def write_excel(path: str, data: dict, sheet_name: str = "Sheet1") -> dict:
    """
    Write data to an Excel file.
    data: { sheet_name: [ {col: val}, ... ] } or flat list of dicts
    """
    try:
        import openpyxl

        p = Path(path).expanduser()

        # Load existing or create new
        if p.exists():
            wb = openpyxl.load_workbook(str(p))
        else:
            wb = openpyxl.Workbook()
            # Remove default empty sheet if we have data
            if "Sheet" in wb.sheetnames:
                del wb["Sheet"]

        # Normalize data format
        if isinstance(data, list):
            data = {sheet_name: data}

        for sname, rows in data.items():
            if sname in wb.sheetnames:
                ws = wb[sname]
            else:
                ws = wb.create_sheet(sname)

            if not rows:
                continue

            # Write headers
            headers = list(rows[0].keys())
            ws.append(headers)

            # Write data rows
            for row in rows:
                ws.append([row.get(h) for h in headers])

            # Basic auto-width
            for col in ws.columns:
                max_len = max(
                    len(str(cell.value)) for cell in col if cell.value is not None
                ) if any(cell.value for cell in col) else 10
                ws.column_dimensions[col[0].column_letter].width = min(max_len + 2, 50)

        p.parent.mkdir(parents=True, exist_ok=True)
        wb.save(str(p))
        return {"success": True, "file_path": str(p), "text": f"✅ Excel saved: {p.name}"}

    except Exception as e:
        logger.error(f"write_excel error: {e}")
        return {"success": False, "error": str(e)}


def update_cell(path: str, sheet: str, row: int, col: Union[int, str], value: Any) -> dict:
    """Update a single cell in an Excel file."""
    try:
        import openpyxl

        p = Path(path).expanduser()
        if not p.exists():
            return {"success": False, "error": f"File not found: {path}"}

        wb = openpyxl.load_workbook(str(p))
        if sheet not in wb.sheetnames:
            return {"success": False, "error": f"Sheet not found: {sheet}"}

        ws = wb[sheet]
        if isinstance(col, str) and col.isalpha():
            ws[f"{col.upper()}{row}"] = value
        else:
            ws.cell(row=int(row), column=int(col)).value = value

        wb.save(str(p))
        return {"success": True, "text": f"✅ Updated {sheet}[{row},{col}] = {value}"}

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── PDF Operations ──────────────────────────────────────────────────────────

def read_pdf(path: str) -> dict:
    """Extract all text from a PDF file."""
    try:
        p = Path(path).expanduser()
        if not p.exists():
            return {"success": False, "error": f"File not found: {path}"}

        text = ""
        try:
            import pdfplumber
            with pdfplumber.open(str(p)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except ImportError:
            import PyPDF2
            with open(str(p), "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"

        if not text.strip():
            return {"success": False, "error": "No text found in PDF (may be scanned/image-based)"}

        return {
            "success": True,
            "content": text[:10000],
            "text": text[:10000],
            "page_count": None,
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


def create_pdf(content: Union[str, dict], output_path: Optional[str] = None) -> dict:
    """
    Create a PDF from text content or structured dict.
    Uses reportlab if available, falls back to fpdf2.
    """
    if not output_path:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = str(_temp_dir() / f"document_{ts}.pdf")

    try:
        # Try reportlab first (better quality)
        _create_pdf_reportlab(content, output_path)
        return {
            "success": True,
            "file_path": output_path,
            "caption": f"Document: {Path(output_path).name}",
            "text": f"✅ PDF created: {output_path}",
        }
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"reportlab failed: {e}, trying fpdf2")

    try:
        _create_pdf_fpdf(content, output_path)
        return {
            "success": True,
            "file_path": output_path,
            "caption": f"Document: {Path(output_path).name}",
            "text": f"✅ PDF created: {output_path}",
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


def _create_pdf_reportlab(content: Union[str, dict], output_path: str):
    """Create PDF using reportlab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=20*mm,
        leftMargin=20*mm,
        topMargin=20*mm,
        bottomMargin=20*mm,
    )

    styles = getSampleStyleSheet()
    story = []

    # Custom styles
    title_style = ParagraphStyle(
        "CustomTitle",
        parent=styles["Title"],
        fontSize=22,
        spaceAfter=6,
        textColor=colors.HexColor("#1a1a2e"),
    )
    h2_style = ParagraphStyle(
        "CustomH2",
        parent=styles["Heading2"],
        fontSize=14,
        spaceBefore=12,
        spaceAfter=4,
        textColor=colors.HexColor("#16213e"),
    )
    body_style = ParagraphStyle(
        "CustomBody",
        parent=styles["BodyText"],
        fontSize=10,
        spaceAfter=6,
        leading=14,
    )
    metric_style = ParagraphStyle(
        "Metric",
        parent=styles["BodyText"],
        fontSize=11,
        textColor=colors.HexColor("#0f3460"),
        spaceAfter=4,
    )

    if isinstance(content, str):
        # Simple text PDF
        for para in content.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip().replace("\n", "<br/>"), body_style))
                story.append(Spacer(1, 4*mm))

    elif isinstance(content, dict):
        # Structured report
        title = content.get("title", "Report")
        date = content.get("date", datetime.now().strftime("%B %d, %Y"))
        summary = content.get("summary", "")
        sections = content.get("sections", [])
        key_metrics = content.get("key_metrics", [])
        conclusion = content.get("conclusion", "")

        # Title block
        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Generated: {date}", styles["Normal"]))
        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#1a1a2e")))
        story.append(Spacer(1, 6*mm))

        # Summary
        if summary:
            story.append(Paragraph("Executive Summary", h2_style))
            story.append(Paragraph(summary, body_style))
            story.append(Spacer(1, 4*mm))

        # Key Metrics
        if key_metrics:
            story.append(Paragraph("Key Metrics", h2_style))
            metric_data = [[m.get("label", ""), str(m.get("value", ""))] for m in key_metrics]
            t = Table(metric_data, colWidths=[100*mm, 60*mm])
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 10),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f0f4f8"), colors.white]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("PADDING", (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 6*mm))

        # Sections
        for section in sections:
            heading = section.get("heading", "")
            sec_content = section.get("content", "")
            table_data = section.get("table")

            if heading:
                story.append(Paragraph(heading, h2_style))

            if sec_content:
                story.append(Paragraph(sec_content.replace("\n", "<br/>"), body_style))
                story.append(Spacer(1, 3*mm))

            if table_data and len(table_data) > 0:
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#16213e")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#f8f9fa"), colors.white]),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("PADDING", (0, 0), (-1, -1), 5),
                    ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ]))
                story.append(t)
                story.append(Spacer(1, 4*mm))

        # Conclusion
        if conclusion:
            story.append(HRFlowable(width="100%", thickness=1, color=colors.grey))
            story.append(Spacer(1, 3*mm))
            story.append(Paragraph("Conclusion & Recommendations", h2_style))
            story.append(Paragraph(conclusion, body_style))

        # Footer
        story.append(Spacer(1, 10*mm))
        story.append(Paragraph(
            f"<i>Generated by GhostPC • {datetime.now().strftime('%Y-%m-%d %H:%M')}</i>",
            ParagraphStyle("Footer", fontSize=8, textColor=colors.grey, alignment=TA_CENTER)
        ))

    doc.build(story)


def _create_pdf_fpdf(content: Union[str, dict], output_path: str):
    """Create PDF using fpdf2 as fallback."""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)

    if isinstance(content, str):
        for line in content.split("\n"):
            pdf.multi_cell(0, 8, line)
    elif isinstance(content, dict):
        title = content.get("title", "Report")
        pdf.set_font("Helvetica", "B", 18)
        pdf.cell(0, 10, title, ln=True, align="C")
        pdf.ln(5)

        pdf.set_font("Helvetica", size=10)
        date = content.get("date", datetime.now().strftime("%B %d, %Y"))
        pdf.cell(0, 8, f"Date: {date}", ln=True)
        pdf.ln(3)

        summary = content.get("summary", "")
        if summary:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Summary", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, summary)
            pdf.ln(4)

        for section in content.get("sections", []):
            heading = section.get("heading", "")
            sec_content = section.get("content", "")
            if heading:
                pdf.set_font("Helvetica", "B", 12)
                pdf.cell(0, 8, heading, ln=True)
            if sec_content:
                pdf.set_font("Helvetica", size=10)
                pdf.multi_cell(0, 6, sec_content)
                pdf.ln(3)

        conclusion = content.get("conclusion", "")
        if conclusion:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Conclusion", ln=True)
            pdf.set_font("Helvetica", size=10)
            pdf.multi_cell(0, 6, conclusion)

    pdf.output(output_path)


def merge_pdfs(paths: list, output_path: Optional[str] = None) -> dict:
    """Merge multiple PDF files into one."""
    try:
        import PyPDF2

        if not output_path:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(_temp_dir() / f"merged_{ts}.pdf")

        writer = PyPDF2.PdfWriter()
        for p in paths:
            reader = PyPDF2.PdfReader(str(Path(p).expanduser()))
            for page in reader.pages:
                writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

        return {
            "success": True,
            "file_path": output_path,
            "caption": f"Merged PDF ({len(paths)} files)",
            "text": f"✅ Merged {len(paths)} PDFs → {output_path}",
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── AI-Powered Report Generation ───────────────────────────────────────────

def generate_report(
    data: Any,
    report_type: str = "summary",
    output_format: str = "pdf",
    output_path: Optional[str] = None,
) -> dict:
    """
    The flagship feature: takes raw data (from Excel or any source),
    uses AI to structure it into a professional report,
    then renders to PDF or DOCX.

    data: dict/list of data (from read_excel output)
    report_type: "production", "sales", "financial", "summary", "custom"
    output_format: "pdf" or "docx"
    """
    try:
        from core.ai import get_ai
        ai = get_ai()

        # Step 1: AI structures the data into a report
        logger.info(f"Generating {report_type} report (AI structuring)...")
        report_structure = ai.generate_report_structure(data, report_type)

        # Step 2: Render to the desired format
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")

        if output_format.lower() == "docx":
            if not output_path:
                output_path = str(_temp_dir() / f"{report_type}_report_{ts}.docx")
            result = create_docx(report_structure, output_path)
        else:
            if not output_path:
                output_path = str(_temp_dir() / f"{report_type}_report_{ts}.pdf")
            result = create_pdf(report_structure, output_path)

        if result.get("success"):
            result["caption"] = f"📊 {report_type.title()} Report"
            result["report_type"] = report_type
            result["output_format"] = output_format
        return result

    except Exception as e:
        logger.error(f"generate_report error: {e}")
        return {"success": False, "error": str(e)}


# ─── DOCX Generation ─────────────────────────────────────────────────────────

def create_docx(content: Union[str, dict], output_path: Optional[str] = None) -> dict:
    """Create a DOCX file from text or structured report dict."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if not output_path:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(_temp_dir() / f"document_{ts}.docx")

        doc = Document()

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

        if isinstance(content, str):
            for para in content.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())

        elif isinstance(content, dict):
            title = content.get("title", "Report")
            date = content.get("date", datetime.now().strftime("%B %d, %Y"))
            summary = content.get("summary", "")
            sections = content.get("sections", [])
            key_metrics = content.get("key_metrics", [])
            conclusion = content.get("conclusion", "")

            # Title
            title_para = doc.add_heading(title, level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph(f"Generated: {date}")
            doc.add_paragraph("─" * 60)

            # Summary
            if summary:
                doc.add_heading("Executive Summary", level=1)
                doc.add_paragraph(summary)

            # Key Metrics table
            if key_metrics:
                doc.add_heading("Key Metrics", level=1)
                table = doc.add_table(rows=1, cols=2)
                table.style = "Table Grid"
                hdr = table.rows[0].cells
                hdr[0].text = "Metric"
                hdr[1].text = "Value"
                for m in key_metrics:
                    row = table.add_row().cells
                    row[0].text = str(m.get("label", ""))
                    row[1].text = str(m.get("value", ""))

            # Sections
            for section in sections:
                heading = section.get("heading", "")
                sec_content = section.get("content", "")
                table_data = section.get("table")

                if heading:
                    doc.add_heading(heading, level=2)
                if sec_content:
                    doc.add_paragraph(sec_content)
                if table_data and len(table_data) > 1:
                    table = doc.add_table(rows=1, cols=len(table_data[0]))
                    table.style = "Table Grid"
                    for i, header in enumerate(table_data[0]):
                        table.rows[0].cells[i].text = str(header)
                    for row_data in table_data[1:]:
                        row = table.add_row().cells
                        for i, val in enumerate(row_data):
                            if i < len(row):
                                row[i].text = str(val)

            # Conclusion
            if conclusion:
                doc.add_heading("Conclusion & Recommendations", level=1)
                doc.add_paragraph(conclusion)

            # Footer
            doc.add_paragraph("")
            footer_para = doc.add_paragraph(
                f"Generated by GhostPC • {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            )
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(output_path)
        return {
            "success": True,
            "file_path": output_path,
            "caption": f"Document: {Path(output_path).name}",
            "text": f"✅ DOCX created: {output_path}",
        }

    except Exception as e:
        logger.error(f"create_docx error: {e}")
        return {"success": False, "error": str(e)}


# ─── Form Filling ────────────────────────────────────────────────────────────

def fill_form(template_path: str, data: dict, output_path: Optional[str] = None) -> dict:
    """
    Fill a DOCX template with {placeholder} values.
    data: { "placeholder": "value" }
    """
    try:
        from docx import Document

        p = Path(template_path).expanduser()
        if not p.exists():
            return {"success": False, "error": f"Template not found: {template_path}"}

        if not output_path:
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(_temp_dir() / f"filled_{p.stem}_{ts}{p.suffix}")

        doc = Document(str(p))

        def replace_in_para(para):
            for key, value in data.items():
                placeholder = f"{{{key}}}"
                if placeholder in para.text:
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        for para in doc.paragraphs:
            replace_in_para(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_para(para)

        doc.save(output_path)
        return {
            "success": True,
            "file_path": output_path,
            "caption": f"Filled form: {Path(output_path).name}",
            "text": f"✅ Form filled: {output_path}",
        }

    except Exception as e:
        return {"success": False, "error": str(e)}


# ─── Excel → Report Pipeline (convenience wrapper) ───────────────────────────

def excel_to_report(
    excel_path: str,
    report_type: str = "summary",
    output_format: str = "pdf",
) -> dict:
    """
    Full pipeline: read Excel → AI report → PDF/DOCX.
    This is the one-command demo moment.
    """
    # Step 1: Read the Excel file
    excel_data = read_excel(excel_path)
    if not excel_data.get("success"):
        return excel_data

    # Step 2: Generate the report
    return generate_report(
        data=excel_data["data"],
        report_type=report_type,
        output_format=output_format,
    )


# ─── Google Sheets (API backend — no browser) ─────────────────────────────────

def _get_gspread_client():
    """
    Return an authenticated gspread client.
    Auth priority:
      1. Service account JSON file at GOOGLE_SHEETS_CREDS_PATH
      2. OAuth2 token cache at ~/.ghostdesk/google_token.json
    Raises ImportError if gspread/google-auth not installed.
    Raises FileNotFoundError if no credentials found.
    """
    import gspread
    from config import USER_DATA_DIR

    # ── Service account ──────────────────────────────────────────────────────
    creds_path_env = os.environ.get("GOOGLE_SHEETS_CREDS_PATH", "")
    default_sa = USER_DATA_DIR / "google_service_account.json"

    if creds_path_env and Path(creds_path_env).exists():
        return gspread.service_account(filename=creds_path_env)
    if default_sa.exists():
        return gspread.service_account(filename=str(default_sa))

    # ── OAuth2 (user browser flow, credentials cached) ───────────────────────
    from google.oauth2.credentials import Credentials
    from google_auth_oauthlib.flow import InstalledAppFlow
    from google.auth.transport.requests import Request

    SCOPES = ["https://www.googleapis.com/auth/spreadsheets"]
    token_path = USER_DATA_DIR / "google_token.json"
    oauth_secret = USER_DATA_DIR / "google_oauth_secret.json"

    creds = None
    if token_path.exists():
        creds = Credentials.from_authorized_user_file(str(token_path), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        elif oauth_secret.exists():
            flow = InstalledAppFlow.from_client_secrets_file(str(oauth_secret), SCOPES)
            creds = flow.run_local_server(port=0)
        else:
            raise FileNotFoundError(
                "No Google credentials found.\n"
                "Option A — Service account: put google_service_account.json in ~/.ghostdesk/\n"
                "Option B — OAuth2: put google_oauth_secret.json in ~/.ghostdesk/\n"
                "Say `how do I set up Google Sheets?` for a step-by-step guide."
            )
        token_path.write_text(creds.to_json())

    return gspread.authorize(creds)


def read_google_sheet(
    url_or_id: str,
    sheet_name: Optional[str] = None,
    range_: Optional[str] = None,
) -> dict:
    """
    Read data from a Google Sheet using the Sheets API — no browser required.

    Args:
        url_or_id: Full Google Sheets URL or just the spreadsheet ID
        sheet_name: Worksheet name (default: first sheet)
        range_: Optional A1-notation range e.g. "A1:D20"
    """
    try:
        import gspread
    except ImportError:
        return {
            "success": False,
            "error": "Install required: pip install gspread google-auth google-auth-oauthlib",
        }

    try:
        gc = _get_gspread_client()

        # Accept both full URL and bare spreadsheet ID
        if "docs.google.com" in url_or_id:
            sh = gc.open_by_url(url_or_id)
        else:
            sh = gc.open_by_key(url_or_id)

        ws = sh.worksheet(sheet_name) if sheet_name else sh.get_worksheet(0)

        if range_:
            values = ws.get(range_)
            rows = [dict(zip(values[0], row)) for row in values[1:]] if len(values) > 1 else []
        else:
            rows = ws.get_all_records()

        title = sh.title
        ws_name = ws.title
        return {
            "success": True,
            "data": rows,
            "total_rows": len(rows),
            "text": (
                f"📊 Google Sheet: *{title}* / *{ws_name}*\n"
                f"Rows: {len(rows)}"
            ),
            "sheet_title": title,
            "worksheet": ws_name,
        }

    except FileNotFoundError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        logger.error(f"read_google_sheet error: {e}")
        return {"success": False, "error": str(e)}


def write_google_sheet(
    url_or_id: str,
    data: list,
    sheet_name: Optional[str] = None,
    start_cell: str = "A1",
    append: bool = False,
) -> dict:
    """
    Write rows to a Google Sheet using the Sheets API — no browser required.

    Args:
        url_or_id: Full Google Sheets URL or spreadsheet ID
        data: List of dicts (headers auto-detected from first dict keys)
              or list of lists (raw rows)
        sheet_name: Worksheet name (default: first sheet)
        start_cell: Top-left cell for writing (default A1)
        append: If True, append rows below existing data instead of overwriting
    """
    try:
        import gspread
    except ImportError:
        return {
            "success": False,
            "error": "Install required: pip install gspread google-auth google-auth-oauthlib",
        }

    try:
        gc = _get_gspread_client()

        if "docs.google.com" in url_or_id:
            sh = gc.open_by_url(url_or_id)
        else:
            sh = gc.open_by_key(url_or_id)

        ws = sh.worksheet(sheet_name) if sheet_name else sh.get_worksheet(0)

        # Normalize: list of dicts → list of lists with header row
        if data and isinstance(data[0], dict):
            headers = list(data[0].keys())
            rows = [headers] + [[row.get(h, "") for h in headers] for row in data]
        else:
            rows = data

        if append:
            ws.append_rows(rows, value_input_option="USER_ENTERED")
            action = "appended"
        else:
            ws.update(start_cell, rows, value_input_option="USER_ENTERED")
            action = "written"

        return {
            "success": True,
            "rows_written": len(rows),
            "text": (
                f"✅ Google Sheet updated: *{sh.title}* / *{ws.title}*\n"
                f"{len(rows)} row(s) {action}."
            ),
        }

    except FileNotFoundError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        logger.error(f"write_google_sheet error: {e}")
        return {"success": False, "error": str(e)}


def update_google_cell(
    url_or_id: str,
    cell: str,
    value: str,
    sheet_name: Optional[str] = None,
) -> dict:
    """Update a single cell in a Google Sheet. cell e.g. 'B3'."""
    try:
        import gspread
    except ImportError:
        return {"success": False, "error": "Install: pip install gspread google-auth"}

    try:
        gc = _get_gspread_client()
        if "docs.google.com" in url_or_id:
            sh = gc.open_by_url(url_or_id)
        else:
            sh = gc.open_by_key(url_or_id)

        ws = sh.worksheet(sheet_name) if sheet_name else sh.get_worksheet(0)
        ws.update_acell(cell, value)

        return {
            "success": True,
            "text": f"✅ Cell *{cell}* updated to `{value}` in *{sh.title}*.",
        }
    except FileNotFoundError as e:
        return {"success": False, "error": str(e)}
    except Exception as e:
        logger.error(f"update_google_cell error: {e}")
        return {"success": False, "error": str(e)}
